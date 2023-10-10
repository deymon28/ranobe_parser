import os
from docx import Document
from docx.opc.exceptions import PackageNotFoundError

def modify_docx_files(folder_path):
    for root, _, files in os.walk(folder_path):
        for file_name in files:
            if file_name.endswith('.docx'):
                file_path = os.path.join(root, file_name)
                try:
                    document = Document(file_path)

                    # Delete the last line
                    paragraphs = document.paragraphs
                    if paragraphs:
                        last_paragraph = paragraphs[-2]
                        document._body._body.remove(last_paragraph._element)

                    # Modify the file name
                    modified_file_name = file_name[:-4].replace('_', ' ')

                    # Add the modified file name at the beginning
                    first_paragraph = document.paragraphs[0]
                    run = first_paragraph.insert_paragraph_before(modified_file_name).runs[0]
                    run.bold = True

                    # Save the modified document
                    document.save(file_path)

                    print("Modified file: " + file_name)
                except PackageNotFoundError:
                    print("Failed to open file: " + file_name)


folder_path = r"C:\Users\dimag\Desktop\Книги\Как держать дистанцию от красавицы"

# Call the function to modify the .docx files
modify_docx_files(folder_path)
