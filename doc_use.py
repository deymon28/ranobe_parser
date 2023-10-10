import os
from docx import Document

folder_path = 'D:\Python_projects\Htmlparse'
for filename in os.listdir(folder_path):
    if filename.endswith('.docx'):
        print(filename)
        file_path = os.path.join(folder_path, filename)
        doc = Document(file_path)
        
        # Delete last line
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.clear()
        
        # Add file name to the beginning of the file
        new_filename = filename[33:].replace('_', ' ')
        doc.paragraphs.insert(0, new_filename)
        
        # Save the modified document
        doc.save(file_path)
